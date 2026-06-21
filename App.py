"""RAG Q&A App.

A Retrieval-Augmented Generation question-answering app built with Streamlit,
FAISS, LangChain (LCEL), HuggingFace embeddings and Groq for fast LLM inference.

Point it at a web page, PDF, Word document, text file or pasted text, then ask
questions about the content. Answers are grounded in the retrieved passages and
the sources used for each answer are shown alongside it.
"""

from __future__ import annotations

import os
from io import BytesIO

import streamlit as st
from dotenv import load_dotenv
from docx import Document as DocxDocument
from pypdf import PdfReader

from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Load variables from a local .env file if present (no-op in production).
load_dotenv()
# WebBaseLoader emits a warning when no user agent is set.
os.environ.setdefault("USER_AGENT", "rag-qa-app/1.0")

# --- Configuration -----------------------------------------------------------
# Embedding model used to vectorise the documents (runs locally on CPU).
EMBEDDING_MODEL = "sentence-transformers/all-mpnet-base-v2"
# Groq chat model. The previous default (deepseek-r1-distill-llama-70b) and
# llama-3.3-70b-versatile have both been deprecated by Groq; gpt-oss-120b is the
# current recommended replacement. Override with the GROQ_MODEL env var.
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
TOP_K = 4

SYSTEM_PROMPT = (
    "You are a helpful assistant answering questions about the provided "
    "documents. Use only the context below to answer. If the answer is not "
    "contained in the context, say you don't know rather than guessing.\n\n"
    "Context:\n{context}"
)


def get_groq_api_key() -> str | None:
    """Return the Groq API key from the environment or Streamlit secrets."""
    key = os.getenv("GROQ_API_KEY")
    if key:
        return key
    try:
        return st.secrets["GROQ_API_KEY"]
    except Exception:
        # No secrets file / key configured — fall back to "no key".
        return None


@st.cache_resource(show_spinner=False)
def load_embeddings() -> HuggingFaceEmbeddings:
    """Load and cache the sentence-transformers embedding model."""
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def _as_bytesio(file_like) -> BytesIO:
    """Normalise an uploaded file (or BytesIO) into a fresh BytesIO buffer."""
    if isinstance(file_like, BytesIO):
        file_like.seek(0)
        return file_like
    if hasattr(file_like, "read"):
        return BytesIO(file_like.read())
    raise ValueError("Expected an uploaded file.")


def _read_pdf(file_like) -> str:
    reader = PdfReader(_as_bytesio(file_like))
    # extract_text() can return None for pages without extractable text.
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(file_like) -> str:
    doc = DocxDocument(_as_bytesio(file_like))
    return "\n".join(para.text for para in doc.paragraphs)


def load_documents(input_type: str, input_data) -> list[Document]:
    """Turn a raw input into LangChain ``Document`` objects with source metadata."""
    if input_type == "Link":
        urls = [u.strip() for u in input_data if u and u.strip()]
        if not urls:
            raise ValueError("Please enter at least one URL.")
        return WebBaseLoader(urls).load()

    if input_data is None:
        raise ValueError("Please provide some input before processing.")

    if input_type == "PDF":
        text, source = _read_pdf(input_data), getattr(input_data, "name", "document.pdf")
    elif input_type == "DOCX":
        text, source = _read_docx(input_data), getattr(input_data, "name", "document.docx")
    elif input_type == "TXT":
        text = _as_bytesio(input_data).read().decode("utf-8")
        source = getattr(input_data, "name", "document.txt")
    elif input_type == "Text":
        if not isinstance(input_data, str) or not input_data.strip():
            raise ValueError("Please enter some text.")
        text, source = input_data, "pasted text"
    else:
        raise ValueError(f"Unsupported input type: {input_type}")

    if not text.strip():
        raise ValueError("No text could be extracted from the input.")
    return [Document(page_content=text, metadata={"source": source})]


def build_vectorstore(documents: list[Document]) -> FAISS:
    """Split documents into chunks and index them in a FAISS vector store."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    chunks = splitter.split_documents(documents)
    if not chunks:
        raise ValueError("The input produced no text to index.")
    return FAISS.from_documents(chunks, load_embeddings())


def _format_docs(docs: list[Document]) -> str:
    return "\n\n".join(doc.page_content for doc in docs)


def answer_question(vectorstore: FAISS, query: str) -> dict:
    """Run a retrieval-augmented generation chain and return the answer + sources.

    Built with LCEL (``langchain-core`` primitives) so it stays compatible across
    LangChain versions, unlike the deprecated ``RetrievalQA`` chain.
    """
    retriever = vectorstore.as_retriever(search_kwargs={"k": TOP_K})
    llm = ChatGroq(model=GROQ_MODEL, temperature=0)
    prompt = ChatPromptTemplate.from_messages(
        [("system", SYSTEM_PROMPT), ("human", "{input}")]
    )

    generate_answer = (
        RunnablePassthrough.assign(context=lambda x: _format_docs(x["docs"]))
        | prompt
        | llm
        | StrOutputParser()
    )
    rag_chain = RunnablePassthrough.assign(
        docs=lambda x: retriever.invoke(x["input"])
    ).assign(answer=generate_answer)

    return rag_chain.invoke({"input": query})


def collect_input(input_type: str):
    """Render the right input widget(s) for the selected input type."""
    if input_type == "Link":
        count = st.number_input(
            "Number of links", min_value=1, max_value=20, value=1, step=1
        )
        return [st.text_input(f"URL {i + 1}", key=f"url_{i}") for i in range(int(count))]
    if input_type == "Text":
        return st.text_area("Paste your text", height=200)
    if input_type == "PDF":
        return st.file_uploader("Upload a PDF file", type=["pdf"])
    if input_type == "TXT":
        return st.file_uploader("Upload a text file", type=["txt"])
    if input_type == "DOCX":
        return st.file_uploader("Upload a Word file", type=["docx", "doc"])
    return None


def main() -> None:
    st.set_page_config(page_title="RAG Q&A App", page_icon="📚")
    st.title("📚 RAG Q&A App")
    st.caption("Ask questions about your PDFs, Word docs, text files or web pages.")

    api_key = get_groq_api_key()
    if api_key:
        os.environ["GROQ_API_KEY"] = api_key
    else:
        st.warning(
            "No Groq API key found. Add `GROQ_API_KEY` to a `.env` file or to "
            "Streamlit secrets to generate answers. You can still index data below."
        )

    with st.sidebar:
        st.header("1 · Add your data")
        input_type = st.selectbox("Input type", ["Link", "PDF", "Text", "DOCX", "TXT"])
        input_data = collect_input(input_type)
        process = st.button("Process", type="primary", use_container_width=True)

    if process:
        try:
            with st.spinner("Reading and indexing your data…"):
                documents = load_documents(input_type, input_data)
                st.session_state.vectorstore = build_vectorstore(documents)
            st.success("Indexed your data. Ask a question below.")
        except Exception as exc:  # surfaced to the user instead of crashing
            st.error(f"Could not process the input: {exc}")

    if "vectorstore" not in st.session_state:
        st.info("👈 Add a data source in the sidebar and click **Process** to begin.")
        return

    st.header("2 · Ask a question")
    query = st.text_input("Your question")
    if st.button("Submit", type="primary") and query:
        if not api_key:
            st.error("Add a Groq API key to generate answers.")
            return
        try:
            with st.spinner("Thinking…"):
                result = answer_question(st.session_state.vectorstore, query)
            st.markdown("### Answer")
            st.write(result["answer"])
            with st.expander("Sources"):
                for i, doc in enumerate(result.get("docs", []), start=1):
                    source = doc.metadata.get("source", "unknown")
                    st.markdown(f"**{i}. {source}**")
                    snippet = doc.page_content.strip()
                    st.write(snippet[:500] + ("…" if len(snippet) > 500 else ""))
        except Exception as exc:
            st.error(f"Could not generate an answer: {exc}")


if __name__ == "__main__":
    main()
